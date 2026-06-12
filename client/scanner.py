"""客户端目录扫描模块"""

import json
import os
import tempfile
import zipfile
from pathlib import Path, PurePosixPath
from typing import Iterable, Optional

import chardet
from loguru import logger

try:
    from docx import Document
except Exception:
    Document = None

try:
    import openpyxl
except Exception:
    openpyxl = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    from local_db import LocalDB
    from matcher import (
        compute_detection_status,
        compute_score,
        compute_sha256,
        compute_simhash,
        match_combined,
        match_keyword,
        match_regex,
        match_semantic_labels,
        match_sha256,
        match_simhash,
        score_to_risk,
    )
except ImportError:
    from client.local_db import LocalDB
    from client.matcher import (
        compute_detection_status,
        compute_score,
        compute_sha256,
        compute_simhash,
        match_combined,
        match_keyword,
        match_regex,
        match_semantic_labels,
        match_sha256,
        match_simhash,
        score_to_risk,
    )

TEXT_SUFFIXES = {
    ".txt", ".csv", ".json", ".xml", ".md", ".py", ".java", ".go", ".sql",
    ".conf", ".config", ".ini", ".yaml", ".yml", ".log", ".properties",
}
OFFICE_SUFFIXES = {".docx", ".xlsx", ".pdf"}
ARCHIVE_SUFFIXES = {".zip"}
UNSUPPORTED_SUFFIXES = {
    ".ppt", ".pptx", ".doc", ".xls", ".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tiff",
    ".eml", ".msg",
}
SKIP_DIRS = {".git", ".hg", ".svn", ".venv", "venv", "node_modules", "__pycache__", "dist", "build", "target", "out", ".mypy_cache", ".pytest_cache", ".idea", ".vscode"}
DEFAULT_MAX_FILE_SIZE = 50 * 1024 * 1024
MAX_FILE_SIZE = int(os.getenv("SCANNER_MAX_FILE_SIZE", DEFAULT_MAX_FILE_SIZE))
MAX_ZIP_DEPTH = int(os.getenv("SCANNER_MAX_ZIP_DEPTH", "2"))
MAX_ZIP_ENTRIES = int(os.getenv("SCANNER_MAX_ZIP_ENTRIES", "200"))
MAX_ZIP_TOTAL_SIZE = int(os.getenv("SCANNER_MAX_ZIP_TOTAL_SIZE", str(100 * 1024 * 1024)))


class ScanResult:
    def __init__(self, file_path: str, file_hash: str, sensitive: bool, sensitive_type: Optional[str], risk_level: str, sensitive_file_id: Optional[str], match_score: int, confidence_level: str, match_detail: dict):
        self.file_path = file_path
        self.file_hash = file_hash
        self.sensitive = sensitive
        self.sensitive_type = sensitive_type
        self.risk_level = risk_level
        self.sensitive_file_id = sensitive_file_id
        self.match_score = match_score
        self.confidence_level = confidence_level
        self.match_detail = match_detail

    def to_dict(self) -> dict:
        return {
            "file_path": self.file_path,
            "file_hash": self.file_hash,
            "sensitive": self.sensitive,
            "sensitive_type": self.sensitive_type,
            "risk_level": self.risk_level,
            "sensitive_file_id": self.sensitive_file_id,
            "match_score": self.match_score,
            "confidence_level": self.confidence_level,
            "match_detail": self.match_detail,
        }


def scan_directory(path: str, db: LocalDB) -> list[dict]:
    root = Path(path)
    if not root.exists():
        raise FileNotFoundError(f"扫描路径不存在: {path}")

    rules = db.load_rules()
    fingerprints = db.load_fingerprints()
    semantic_labels = db.load_semantic_labels()
    config = db.load_config()
    simhash_threshold = int(config.get("simhash_threshold", 3))
    logger.info(f"加载本地规则: {len(rules)} 条, 指纹: {len(fingerprints)} 条, 语义标签: {len(semantic_labels)} 条, SimHash阈值: {simhash_threshold}")

    results = []
    for file_path in iter_files(root):
        try:
            if file_path.suffix.lower() in ARCHIVE_SUFFIXES:
                zip_results = scan_zip_file(file_path, rules, fingerprints, semantic_labels, simhash_threshold=simhash_threshold)
                for result in zip_results:
                    db.upsert_file_tag(
                        file_path=result.file_path,
                        file_hash=result.file_hash,
                        sensitive=result.sensitive,
                        sensitive_type=result.sensitive_type,
                        risk_level=result.risk_level,
                        sensitive_file_id=result.sensitive_file_id,
                        match_score=result.match_score,
                        confidence_level=result.confidence_level,
                        match_detail=result.match_detail,
                    )
                    results.append(result.to_dict())
                    if result.confidence_level != "clean":
                        logger.warning(f"发现命中文件: {result.file_path}, confidence={result.confidence_level}, score={result.match_score}, risk={result.risk_level}")
                    else:
                        logger.info(f"扫描完成: {result.file_path}, score={result.match_score}")
                continue
            result = scan_file(file_path, rules, fingerprints, semantic_labels, simhash_threshold=simhash_threshold)
            db.upsert_file_tag(
                file_path=result.file_path,
                file_hash=result.file_hash,
                sensitive=result.sensitive,
                sensitive_type=result.sensitive_type,
                risk_level=result.risk_level,
                sensitive_file_id=result.sensitive_file_id,
                match_score=result.match_score,
                confidence_level=result.confidence_level,
                match_detail=result.match_detail,
            )
            results.append(result.to_dict())
            if result.confidence_level != "clean":
                logger.warning(f"发现命中文件: {result.file_path}, confidence={result.confidence_level}, score={result.match_score}, risk={result.risk_level}")
            else:
                logger.info(f"扫描完成: {result.file_path}, score={result.match_score}")
        except Exception as exc:
            logger.error(f"扫描失败: {file_path}, reason={exc}")
    return results


def iter_files(root: Path) -> Iterable[Path]:
    if root.is_file():
        yield root
        return

    for file_path in root.rglob("*"):
        if any(part in SKIP_DIRS for part in file_path.parts):
            continue
        if not file_path.is_file():
            continue
        if file_path.stat().st_size > MAX_FILE_SIZE:
            logger.warning(f"跳过过大文件: {file_path}")
            continue
        yield file_path


def is_safe_zip_member(name: str) -> bool:
    normalized = name.replace("\\", "/")
    path = PurePosixPath(normalized)
    return not path.is_absolute() and ".." not in path.parts


def scan_zip_file(zip_path: Path, rules: list, fingerprints: list, semantic_labels: Optional[dict] = None, simhash_threshold: int = 3, depth: int = 0) -> list[ScanResult]:
    if depth >= MAX_ZIP_DEPTH:
        logger.warning(f"跳过超过递归层级限制的 ZIP: {zip_path}")
        return []

    results = []
    total_size = 0
    scanned_entries = 0
    try:
        with zipfile.ZipFile(zip_path) as archive:
            for info in archive.infolist():
                if info.is_dir():
                    continue
                if scanned_entries >= MAX_ZIP_ENTRIES:
                    logger.warning(f"ZIP 条目数超过限制，停止扫描: {zip_path}")
                    break
                if not is_safe_zip_member(info.filename):
                    logger.warning(f"跳过存在 Zip Slip 风险的条目: {zip_path}!{info.filename}")
                    continue
                if info.file_size > MAX_FILE_SIZE:
                    logger.warning(f"跳过 ZIP 内过大文件: {zip_path}!{info.filename}")
                    continue
                total_size += info.file_size
                if total_size > MAX_ZIP_TOTAL_SIZE:
                    logger.warning(f"ZIP 解压总量超过限制，停止扫描: {zip_path}")
                    break

                scanned_entries += 1
                member_name = PurePosixPath(info.filename).name or "member"
                virtual_path = f"{zip_path}!{info.filename}"
                with archive.open(info) as member:
                    data = member.read(MAX_FILE_SIZE + 1)
                if len(data) > MAX_FILE_SIZE:
                    logger.warning(f"跳过 ZIP 内过大文件: {virtual_path}")
                    continue
                if member_name.lower().endswith(".zip"):
                    with tempfile.NamedTemporaryFile(suffix=".zip", delete=False) as tmp:
                        tmp.write(data)
                        nested_path = Path(tmp.name)
                    try:
                        for result in scan_zip_file(nested_path, rules, fingerprints, semantic_labels, simhash_threshold, depth + 1):
                            result.file_path = f"{virtual_path}!{result.file_path.split('!', 1)[-1]}"
                            results.append(result)
                    finally:
                        nested_path.unlink(missing_ok=True)
                    continue

                suffix = PurePosixPath(member_name).suffix.lower()
                source_path = None
                if suffix in OFFICE_SUFFIXES:
                    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                        tmp.write(data)
                        source_path = Path(tmp.name)
                try:
                    result = scan_bytes(virtual_path, member_name, data, rules, fingerprints, semantic_labels, simhash_threshold=simhash_threshold, source_path=source_path)
                    results.append(result)
                finally:
                    if source_path is not None:
                        source_path.unlink(missing_ok=True)
    except zipfile.BadZipFile as exc:
        logger.warning(f"ZIP 解析失败: {zip_path}, reason={exc}")
    return results


def scan_bytes(display_path: str, file_name: str, data: bytes, rules: list, fingerprints: list, semantic_labels: Optional[dict] = None, simhash_threshold: int = 3, source_path: Optional[Path] = None) -> ScanResult:
    sha256 = compute_sha256(data)

    sha_hit = match_sha256(sha256, fingerprints)
    semantic_labels = semantic_labels or {}
    if sha_hit:
        label_detail = semantic_labels.get(sha_hit.get("sensitive_file_id"), {})
        detail = {
            "sha256_hit": True,
            "simhash_hit": False,
            "regex_hits": [],
            "keyword_hits": [],
            "combined_hits": [],
            "semantic_label_hits": [],
            "semantic_labels": label_detail.get("semantic_labels", []),
            "embedding_id": label_detail.get("embedding_id"),
        }
        return ScanResult(display_path, sha256, True, "样本指纹命中", "high", sha_hit.get("sensitive_file_id"), 100, "sensitive", detail)

    text = ""
    extract_error = None
    skip_reason = None
    scan_path = source_path or Path(file_name)
    if scan_path.suffix.lower() in UNSUPPORTED_SUFFIXES:
        skip_reason = "unsupported_format"
    else:
        try:
            text = extract_text_from_bytes(scan_path, data, source_path=source_path)
        except Exception as exc:
            extract_error = str(exc)
            logger.warning(f"文本提取失败: {display_path}, reason={extract_error}")
        if not text and scan_path.suffix.lower() == ".pdf":
            skip_reason = "pdf_text_empty" if PdfReader is not None else "pdf_reader_missing"

    simhash = compute_simhash(text) if text else ""
    sim_hit = match_simhash(simhash, fingerprints, threshold=simhash_threshold) if simhash else None
    regex_hits = match_regex(text, rules) if text else []
    keyword_hits = match_keyword(text, rules) if text else []
    combined_hits = match_combined(text, rules) if text else []
    semantic_hits = match_semantic_labels(text, semantic_labels) if text else []

    score = compute_score(bool(sha_hit), bool(sim_hit), regex_hits, keyword_hits, combined_hits, semantic_hits)
    sensitive, confidence_level = compute_detection_status(score)
    risk_level = score_to_risk(score)
    sensitive_type = infer_sensitive_type(regex_hits, keyword_hits, combined_hits, rules)
    if sensitive_type is None and semantic_hits:
        sensitive_type = semantic_hits[0].get("semantic_label") or "语义标签命中"
    sensitive_file_id = sim_hit.get("sensitive_file_id") if sim_hit else None
    if sensitive_file_id is None and semantic_hits:
        sensitive_file_id = semantic_hits[0].get("sensitive_file_id")
    label_detail = semantic_labels.get(sensitive_file_id, {}) if sensitive_file_id else {}

    detail = {
        "sha256_hit": False,
        "simhash_hit": bool(sim_hit),
        "simhash": simhash,
        "regex_hits": regex_hits,
        "keyword_hits": keyword_hits,
        "combined_hits": combined_hits,
        "semantic_label_hits": semantic_hits,
        "semantic_labels": label_detail.get("semantic_labels", []),
        "embedding_id": label_detail.get("embedding_id"),
    }
    if extract_error:
        detail["extract_error"] = extract_error
    if skip_reason:
        detail["skip_reason"] = skip_reason
    return ScanResult(display_path, sha256, sensitive, sensitive_type, risk_level, sensitive_file_id, score, confidence_level, detail)


def scan_file(file_path: Path, rules: list, fingerprints: list, semantic_labels: Optional[dict] = None, simhash_threshold: int = 3) -> ScanResult:
    data = file_path.read_bytes()
    return scan_bytes(str(file_path), file_path.name, data, rules, fingerprints, semantic_labels, simhash_threshold=simhash_threshold, source_path=file_path)


def extract_text(file_path: Path, data: bytes) -> str:
    return extract_text_from_bytes(file_path, data, source_path=file_path)


def extract_text_from_bytes(file_path: Path, data: bytes, source_path: Optional[Path] = None) -> str:
    suffix = file_path.suffix.lower()
    path_for_parser = source_path or file_path
    if suffix in TEXT_SUFFIXES:
        return read_text_with_detection(data)
    if suffix == ".docx":
        return extract_docx(path_for_parser)
    if suffix == ".xlsx":
        return extract_xlsx(path_for_parser)
    if suffix == ".pdf":
        return extract_pdf(path_for_parser)
    return read_text_with_detection(data)


def read_text_with_detection(data: bytes) -> str:
    detection = chardet.detect(data)
    encoding = detection.get("encoding") or "utf-8"
    try:
        return data.decode(encoding, errors="ignore")
    except LookupError:
        return data.decode("utf-8", errors="ignore")


def extract_docx(file_path: Path) -> str:
    if Document is None:
        logger.warning("未安装 python-docx，跳过 docx 文本提取")
        return ""
    doc = Document(str(file_path))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells if cell.text)
    return "\n".join(parts)


def extract_xlsx(file_path: Path) -> str:
    if openpyxl is None:
        logger.warning("未安装 openpyxl，跳过 xlsx 文本提取")
        return ""
    wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            values = [str(cell) for cell in row if cell is not None]
            if values:
                parts.append(" ".join(values))
    wb.close()
    return "\n".join(parts)


def extract_pdf(file_path: Path) -> str:
    if PdfReader is None:
        logger.warning("未安装 pypdf，跳过 pdf 文本提取")
        return ""
    reader = PdfReader(str(file_path))
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text:
            parts.append(text)
    return "\n".join(parts)


def infer_sensitive_type(regex_hits: list, keyword_hits: list, combined_hits: list, rules: list) -> Optional[str]:
    for hit_group in (keyword_hits, combined_hits, regex_hits):
        for hit in hit_group:
            rule_id = hit.get("rule_id")
            for rule in rules:
                if rule.get("rule_id") == rule_id:
                    return rule.get("sensitive_type")
    return None


def dump_results(results: list[dict]):
    print(json.dumps(results, ensure_ascii=False, indent=2))
